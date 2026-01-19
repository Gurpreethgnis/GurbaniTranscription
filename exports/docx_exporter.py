"""
DOCX Exporter for Formatted Documents.

Exports FormattedDocument to Microsoft Word (DOCX) format with
styled sections and Gurbani quote formatting.
"""
import logging
from pathlib import Path
from typing import List, Optional, Any

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    Document = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore
    Inches = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    DOCX_AVAILABLE = False

from core.models import FormattedDocument, DocumentSection, QuoteContent
from exports.base_exporter import BaseExporter
from core.errors import ExportError

logger = logging.getLogger(__name__)


class DOCXExporter(BaseExporter):
    """
    Exports formatted documents to DOCX format.
    
    Produces styled Word documents with:
    - Section headings
    - Formatted Gurbani quotes
    - Metadata
    """
    
    def __init__(self):
        """Initialize DOCX exporter."""
        super().__init__("docx", ".docx")
        
        if not DOCX_AVAILABLE:
            raise ExportError(
                "docx",
                "python-docx is not installed. Install with: pip install python-docx"
            )
    
    def _export_impl(
        self,
        document: FormattedDocument,
        output_path: Path
    ) -> Path:
        """
        Export document to DOCX file.
        
        Args:
            document: FormattedDocument to export
            output_path: Path where DOCX file should be saved
        
        Returns:
            Path to exported DOCX file
        """
        # Create new document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = document.title
        doc.core_properties.author = "KathaTranscription"
        
        # Add title
        title_para = doc.add_heading(document.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Source: {document.source_file}")
        doc.add_paragraph(f"Created: {document.created_at}")
        if document.metadata:
            total = document.metadata.get('total_segments', len(document.sections))
            doc.add_paragraph(f"Total Sections: {total}")
        
        doc.add_paragraph()  # Blank line
        
        # Add sections
        for section in document.sections:
            self._add_section(doc, section)
        
        # Save document
        doc.save(str(output_path))
        
        logger.debug(f"Exported DOCX document: {output_path} ({output_path.stat().st_size} bytes)")
        
        return output_path
    
    def _add_section(self, doc: Any, section: DocumentSection) -> None:
        """
        Add a document section to the Word document.
        
        Args:
            doc: Document object
            section: DocumentSection to add
        """
        # Section header
        header = self._get_section_header(section.section_type)
        if header:
            doc.add_heading(header, level=2)
        
        # Content
        if isinstance(section.content, QuoteContent):
            self._add_quote(doc, section.content)
        else:
            # Regular text content
            text = str(section.content)
            para = doc.add_paragraph(text)
            para.style = 'Normal'
        
        # Timestamp
        if section.start_time is not None:
            time_str = self._format_timestamp(section.start_time)
            para = doc.add_paragraph(f"[Time: {time_str}]")
            para.style = 'Intense Quote'  # Italic style
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_quote(self, doc: Any, quote: QuoteContent) -> None:
        """
        Add a Gurbani quote to the document with formatting.
        
        Args:
            doc: Document object
            quote: QuoteContent to add
        """
        # Gurmukhi text (primary, large, centered, bold)
        gurmukhi_para = doc.add_paragraph()
        gurmukhi_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = gurmukhi_para.add_run(quote.gurmukhi)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = 'Gurmukhi'  # Will use system default if not available
        
        # Roman transliteration (italicized, smaller)
        if quote.roman:
            roman_para = doc.add_paragraph()
            roman_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = roman_para.add_run(quote.roman)
            run.font.size = Pt(12)
            run.font.italic = True
        
        # English translation (if available)
        if quote.english_translation:
            english_para = doc.add_paragraph()
            english_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = english_para.add_run(quote.english_translation)
            run.font.size = Pt(11)
            run.font.italic = True
        
        # Metadata (smaller, centered)
        metadata_parts = []
        if quote.source:
            metadata_parts.append(f"Source: {quote.source}")
        if quote.ang:
            metadata_parts.append(f"Ang: {quote.ang}")
        if quote.raag:
            metadata_parts.append(f"Raag: {quote.raag}")
        if quote.author:
            metadata_parts.append(f"Author: {quote.author}")
        
        if metadata_parts:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta_para.add_run(f"({', '.join(metadata_parts)})")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
        
        # Context lines (if available)
        if quote.context_lines:
            doc.add_paragraph("Context:")
            for context_line in quote.context_lines:
                para = doc.add_paragraph(context_line, style='List Bullet')
    
    def _get_section_header(self, section_type: str) -> Optional[str]:
        """Get section header text."""
        headers = {
            "opening_gurbani": "Opening Gurbani",
            "fateh": "Fateh",
            "topic": "Topic",
            "quote": "Gurbani Quote",
            "katha": None  # No header for regular katha
        }
        return headers.get(section_type)
    
    def _format_timestamp(self, seconds: float) -> str:
        """Format timestamp as MM:SS."""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"
